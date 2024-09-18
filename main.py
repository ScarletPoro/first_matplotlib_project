import os
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import numpy as np
import io

# Imposta il percorso del desktop
desktop_path = r'C:\Users\wepeko\OneDrive\Desktop'

# Verifica e crea la cartella Desktop se non esiste
if not os.path.exists(desktop_path):
    os.makedirs(desktop_path)

# Dati degli utenti in forma di liste
dati_utenti = [
    {'nome': 'Mario', 'eta': 35, 'peso': 85, 'altezza': 1.75},
    {'nome': 'Luigi', 'eta': 42, 'peso': 95, 'altezza': 1.82},
    {'nome': 'Giulia', 'eta': 28, 'peso': 68, 'altezza': 1.65},
    {'nome': 'Carla', 'eta': 50, 'peso': 102, 'altezza': 1.70},
    {'nome': 'Pietro', 'eta': 31, 'peso': 77, 'altezza': 1.78}
]

# Estrazione dei dati per i grafici
nomi = [utente['nome'] for utente in dati_utenti]
eta = [utente['eta'] for utente in dati_utenti]
peso = [utente['peso'] for utente in dati_utenti]
altezza = [utente['altezza'] for utente in dati_utenti]

# === GRAFICO SCATTER (Età vs Peso) ===
plt.figure(figsize=(10, 5))
plt.scatter(eta, peso, color='blue', marker='o')
plt.title('Età vs Peso')
plt.xlabel('Età')
plt.ylabel('Peso')
plt.grid(True)

# Salvataggio del grafico scatter in un buffer
scatter_buffer = io.BytesIO()
plt.savefig(scatter_buffer, format='png')
scatter_buffer.seek(0)

# === GRAFICO A TORTA (Distribuzione del Peso) ===
plt.figure(figsize=(8, 8))
plt.pie(peso, labels=nomi, autopct='%1.1f%%', startangle=140)
plt.title('Distribuzione del Peso tra gli Utenti')

# Salvataggio del grafico a torta in un buffer
pie_buffer = io.BytesIO()
plt.savefig(pie_buffer, format='png')
pie_buffer.seek(0)

# === GRAFICO A STELLA / RADAR (Caratteristiche degli Utenti) ===
# Creare un grafico radar per ogni utente con le caratteristiche: Età, Peso, Altezza
caratteristiche = ['Età', 'Peso', 'Altezza']
num_caratteristiche = len(caratteristiche)

# Normalizzazione dei dati per il grafico radar
dati_radar = []
for utente in dati_utenti:
    dati_radar.append([utente['eta'] / 100, utente['peso'] / 120, utente['altezza'] / 2])

# Impostazioni del grafico radar
angles = np.linspace(0, 2 * np.pi, num_caratteristiche, endpoint=False).tolist()
angles += angles[:1]  # Chiudi il cerchio

plt.figure(figsize=(8, 8))
ax = plt.subplot(111, polar=True)

# Tracciare un radar plot per ogni utente
for i, dati in enumerate(dati_radar):
    dati += dati[:1]  # Chiudi il cerchio
    ax.plot(angles, dati, linewidth=2, linestyle='solid', label=nomi[i])
    ax.fill(angles, dati, alpha=0.25)

# Impostazioni del grafico radar
ax.set_yticklabels([])
ax.set_xticks(angles[:-1])
ax.set_xticklabels(caratteristiche)
plt.title('Caratteristiche degli Utenti')
plt.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1))

# Salvataggio del grafico radar in un buffer
radar_buffer = io.BytesIO()
plt.savefig(radar_buffer, format='png')
radar_buffer.seek(0)

# === CREAZIONE DEL DOCUMENTO WORD ===
doc = Document()
doc.add_heading('Report Utenti', 0)

# Aggiunta di descrizioni e grafici al documento
doc.add_paragraph('Questo documento mostra diversi grafici basati sui dati degli utenti.')

# Aggiungi il grafico scatter al documento
doc.add_heading('Grafico Scatter: Età vs Peso', level=1)
doc.add_paragraph('Il grafico seguente mostra la distribuzione del peso in funzione dell\'età.')
doc.add_picture(scatter_buffer, width=Inches(5.0))

# Aggiungi il grafico a torta al documento
doc.add_heading('Grafico a Torta: Distribuzione del Peso', level=1)
doc.add_paragraph('Il grafico a torta mostra la distribuzione del peso tra gli utenti.')
doc.add_picture(pie_buffer, width=Inches(5.0))

# Aggiungi il grafico radar al documento
doc.add_heading('Grafico Radar: Caratteristiche degli Utenti', level=1)
doc.add_paragraph('Il grafico radar rappresenta le caratteristiche degli utenti: età, peso e altezza.')
doc.add_picture(radar_buffer, width=Inches(5.0))

# Salvataggio del documento Word sul desktop
doc.save(os.path.join(desktop_path, 'report_utenti.docx'))

# Chiudi i buffer delle immagini
scatter_buffer.close()
pie_buffer.close()
radar_buffer.close()

print(f"Documento creato con successo sul Desktop: {desktop_path}")
